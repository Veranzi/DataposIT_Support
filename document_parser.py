import os
import docx


def parse_documents(folder_path):
    text_chunks = []
    image_list = []
    
    print(f"Looking for documents in: {folder_path}")
    
    if not os.path.exists(folder_path):
        print(f"Warning: Folder {folder_path} does not exist!")
        return text_chunks, image_list
    
    files = os.listdir(folder_path)
    print(f"Found {len(files)} files in folder")
    
    for file in files:
        path = os.path.join(folder_path, file)
        print(f"Processing file: {file}")
        
        try:
            if file.endswith(".pdf"):
                print(f"  Skipping PDF: {file} - OCR not available")
                continue

            elif file.endswith(".docx"):
                print(f"  Parsing DOCX: {file}")
                try:
                    doc = docx.Document(path)
                    full_text = ""
                    
                    # Extract text from paragraphs
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            full_text += paragraph.text + "\n"
                    
                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    full_text += cell.text + "\n"
                    
                    # Extract text from headers and footers
                    for section in doc.sections:
                        for header in section.header.paragraphs:
                            if header.text.strip():
                                full_text += header.text + "\n"
                        for footer in section.footer.paragraphs:
                            if footer.text.strip():
                                full_text += footer.text + "\n"
                    
                    if full_text.strip():
                        text_chunks.append(full_text.strip())
                        print(f"    Extracted: {len(full_text)} characters")
                        print(f"    Sample: {full_text[:200]}...")
                    else:
                        print(f"    Warning: No text extracted from {file}")
                        
                except Exception as e:
                    print(f"    Error parsing DOCX {file}: {str(e)}")

            elif file.endswith(".txt"):
                print(f"  Parsing TXT: {file}")
                with open(path, 'r', encoding='utf-8') as f:
                    text = f.read()
                    text_chunks.append(text)
                    print(f"    Extracted: {len(text)} characters")

            elif file.endswith(".doc"):
                print(f"  Skipping DOC file (not supported): {file}")
                continue

            else:
                print(f"  Skipping unsupported file type: {file}")
                continue
                
        except Exception as e:
            print(f"  Error processing {file}: {str(e)}")
            continue
    
    print(f"Total text chunks: {len(text_chunks)}")
    print(f"Total images: {len(image_list)}")
    
    # Print a sample of the extracted text for debugging
    if text_chunks:
        print("Sample of extracted text:")
        for i, chunk in enumerate(text_chunks[:2]):  # Show first 2 chunks
            print(f"Chunk {i+1}: {chunk[:200]}...")
    
    return text_chunks, image_list